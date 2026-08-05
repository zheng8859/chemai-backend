"""Exam export service — Word docx document generation (25号 §十).

排版规格:
  - A4 (210mm x 297mm), margins: top 2.5cm, bottom 2.0cm, left 2.5cm, right 2.0cm
  - Font: SimSun 11pt body, 16pt bold centered title
  - 密封线 (seal line): name / class / score fields
  - 5 question type groupings with numbered sections
  - Dual mode: student (no answers) / teacher (red answers + green analysis)
"""

import io
import logging
from datetime import datetime, timezone
from typing import Optional

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from ..models.teaching import ExamRecord
from ..services.exam_management_service import ExamManagementService

logger = logging.getLogger(__name__)

# Type labels
TYPE_LABELS = {
    "choice": "选择题",
    "fill_blank": "填空题",
    "calculation": "计算题",
    "experiment_inquiry": "实验题",
    "equation_balancing": "推断题",
}

TYPE_ORDER = ["choice", "fill_blank", "calculation", "experiment_inquiry", "equation_balancing"]


class ExamExportError(Exception):
    def __init__(self, detail: str):
        self.detail = detail
        super().__init__(detail)


class ExamExportService:

    @staticmethod
    async def export_to_docx(
        db: AsyncSession,
        exam_id: int,
        with_answers: bool = False,
    ) -> io.BytesIO:
        """Generate Word document for an exam.

        Args:
            db: 数据库会话
            exam_id: 考试记录 ID
            with_answers: True=教师版 (含答案+解析), False=学生版

        Returns:
            io.BytesIO with the .docx file content
        """
        # Load exam record
        result = await db.execute(
            select(ExamRecord).where(ExamRecord.id == exam_id)
        )
        exam = result.scalar_one_or_none()
        if exam is None:
            raise ExamExportError(f"考试不存在: id={exam_id}")

        # Load questions
        questions = await ExamManagementService.get_exam_questions(db, exam_id)
        if not questions:
            raise ExamExportError("该考试暂无题目，无法导出")

        # Group by type
        grouped = {}
        for q in questions:
            t = q.get("question_type", "choice")
            grouped.setdefault(t, []).append(q)

        # Build document
        doc = Document()
        _setup_page(doc)
        _add_seal_line(doc, exam.name or "试卷")

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(exam.name or "化学试卷")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = "SimSun"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

        # Subtitle
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = sub.add_run(f"考试日期: {exam.exam_date}    题目数: {len(questions)}")
        run2.font.size = Pt(10)
        run2.font.name = "SimSun"
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

        doc.add_paragraph()  # spacing

        # Grouped questions
        section_num = 1
        for t in TYPE_ORDER:
            qs = grouped.get(t, [])
            if not qs:
                continue

            # Section header
            header = doc.add_paragraph()
            run_h = header.add_run(_section_name(section_num, t))
            run_h.font.size = Pt(14)
            run_h.font.bold = True
            run_h.font.name = "SimSun"
            run_h._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

            for qi, q in enumerate(qs, 1):
                _add_question(doc, qi, q, with_answers)

            section_num += 1

        # Teacher version footer
        if with_answers:
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_f = footer.add_run("（含答案版）")
            run_f.font.size = Pt(10)
            run_f.font.color.rgb = RGBColor(0xB4, 0x3C, 0x28)
            run_f.font.name = "SimSun"
            run_f._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

        # Write to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf


# ═══════════════════════════════════════════════
# Internal helpers
# ═══════════════════════════════════════════════

def _setup_page(doc: Document):
    """Set A4 page with margins."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)


def _add_seal_line(doc: Document, exam_name: str):
    """Add seal line area with name/class/score fields."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    seal_text = (
        f"姓名：____________    班级：____________    得分：____________"
    )
    run = p.add_run(seal_text)
    run.font.size = Pt(10)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # Separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sep.add_run("—" * 40)
    run_s.font.size = Pt(8)
    run_s.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def _section_name(num: int, qtype: str) -> str:
    labels = {
        "choice": "选择题",
        "fill_blank": "填空题",
        "calculation": "计算题",
        "experiment_inquiry": "实验题",
        "equation_balancing": "推断题",
    }
    chinese = ["", "一", "二", "三", "四", "五"]
    return f"{chinese[num] if num < 6 else num}、{labels.get(qtype, qtype)}"


def _add_question(doc: Document, num: int, q: dict, with_answers: bool):
    """Add a single question to the document."""
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {q.get('content', '')}")
    run.font.size = Pt(11)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # Options (for choice questions)
    options = q.get("options") or []
    if options:
        for opt in options:
            opt_p = doc.add_paragraph()
            opt_p.paragraph_format.left_indent = Cm(1.0)
            run_o = opt_p.add_run(opt)
            run_o.font.size = Pt(11)
            run_o.font.name = "SimSun"
            run_o._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # Answer area for non-choice types
    qtype = q.get("question_type", "choice")
    if qtype in ("calculation", "experiment_inquiry", "equation_balancing"):
        space_p = doc.add_paragraph()
        space_p.paragraph_format.space_before = Pt(12)
        run_sp = space_p.add_run("解：")
        run_sp.font.size = Pt(11)
        run_sp.font.name = "SimSun"
        run_sp._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        # Add blank lines for answer area
        for _ in range(3):
            blank = doc.add_paragraph()
            blank_p = blank.add_run(" ")
            blank_p.font.size = Pt(11)

    # Answer + analysis (teacher version only)
    if with_answers:
        ans = q.get("answer", "")
        if ans:
            ans_p = doc.add_paragraph()
            ans_p.paragraph_format.left_indent = Cm(0.5)
            run_a = ans_p.add_run(f"【答案】{ans}")
            run_a.font.size = Pt(10)
            run_a.font.color.rgb = RGBColor(0xB4, 0x3C, 0x28)  # Error red
            run_a.font.name = "SimSun"
            run_a._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

        analysis = q.get("analysis", "")
        if analysis:
            ana_p = doc.add_paragraph()
            ana_p.paragraph_format.left_indent = Cm(0.5)
            run_an = ana_p.add_run(f"【解析】{analysis}")
            run_an.font.size = Pt(10)
            run_an.font.color.rgb = RGBColor(0x2C, 0x6E, 0x49)  # Success green
            run_an.font.name = "SimSun"
            run_an._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # Spacing between questions
    doc.add_paragraph()
